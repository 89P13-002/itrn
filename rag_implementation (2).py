def main():
    logger.info("Starting RAG model implementation with similarity-based metadata clustering...")
    
    # Initialize components
    llm_client = LLMClient()
    
    # Load embedding models
    content_embedding_model = load_embedding_model(config.EMBEDDING_MODEL_NAME)
    metadata_embedding_model = load_embedding_model("all-MiniLM-L6-v2")  # For metadata similarity
    
    # Initialize hierarchical FAISS index with metadata embedding model
    faiss_index = HierarchicalFAISSIndex(config.EMBEDDING_DIM, metadata_embedding_model)
    
    # Load and preprocess documents
    logger.info("Loading and preprocessing documents...")
    documents = load_and_preprocess_lectures(config.LECTURE_FOLDER)
    
    if not documents:
        logger.error("No documents loaded. Please check the lecture folder.")
        return
    
    # Process each document with LLM chunking and metadata extraction
    all_chunks = []
    all_metadata = []
    
    for doc in documents:
        logger.info(f"Processing document: {doc['filename']}")
        
        # Intelligent chunking using LLM
        chunks = llm_client.intelligent_chunk(doc['content'], config.MAX_CHUNK_SIZE)
        
        for chunk in chunks:
            # Extract metadata using LLM
            metadata = llm_client.extract_metadata(chunk)
            metadata['source_file'] = doc['filename']
            metadata['chunk_text'] = chunk
            
            all_chunks.append(chunk)
            all_metadata.append(metadata)
        
        logger.info(f"Created {len(chunks)} chunks from {doc['filename']}")
    
    logger.info(f"Total chunks created: {len(all_chunks)}")
    
    # Compute embeddings for chunks (content embeddings)
    logger.info("Computing content embeddings for chunks...")
    chunk_embeddings = compute_embeddings(all_chunks, content_embedding_model, config.BATCH_SIZE)
    chunk_embeddings = normalize_embeddings(chunk_embeddings)
    
    # Build hierarchical FAISS index with similarity-based metadata clustering
    logger.info("Building hierarchical FAISS index with metadata similarity clustering...")
    faiss_index.build_hierarchical_index(chunk_embeddings, all_metadata)
    
    # Save the index
    if not os.path.exists(config.INDEX_PATH):
        os.makedirs(config.INDEX_PATH)
    faiss_index.save_index(config.INDEX_PATH)
    
    # Load and process queries
    queries = load_and_preprocess_queries(config.QUERY_FILE)
    
    if queries:
        logger.info("Processing queries...")
        query_embeddings = compute_embeddings(queries, content_embedding_model, config.BATCH_SIZE)
        query_embeddings = normalize_embeddings(query_embeddings)
        
        # Evaluate the index
        metrics = evaluate_faiss_index(faiss_index, query_embeddings)
        
        # Test search functionality with parallel coarse cluster search
        logger.info("Testing search functionality with parallel coarse cluster search...")
        for i, (query, query_emb) in enumerate(zip(queries[:3], query_embeddings[:3])):
            logger.info(f"\nQuery {i+1}: {query}")
            
            # Search with different numbers of coarse clusters
            for num_coarse in [2, 3, 4]:
                logger.info(f"\n  Searching top {num_coarse} coarse clusters:")
                results = faiss_index.search(query_emb, k=5, num_coarse_clusters=num_coarse)
                
                logger.info(f"  Top 5 results from {num_coarse} coarse clusters:")
                for j, (score, idx, metadata) in enumerate(results):
                    logger.info(f"    {j+1}. Score: {score:.4f}")
                    logger.info(f"       Topic: {metadata.get('topic', 'N/A')}")
                    logger.info(f"       Category: {metadata.get('category', 'N/A')}")
                    logger.info(f"       Summary: {metadata.get('summary', 'N/A')[:100]}...")
                    logger.info(f"       Source: {metadata.get('source_file', 'N/A')}")
        
        # Demonstrate cluster analysis
        logger.info("\n" + "="*50)
        logger.info("CLUSTER ANALYSIS")
        logger.info("="*50)
        
        logger.info(f"Total coarse clusters: {len(faiss_index.coarse_clusters)}")
        for cluster_id, cluster_info in faiss_index.coarse_clusters.items():
            logger.info(f"\nCoarse Cluster: {cluster_id}")
            logger.info(f"  Size: {len(cluster_info['indices'])} chunks")
            logger.info(f"  Summary: {cluster_info['centroid_text']}")
            
            # Show sample metadata from this cluster
            sample_metadata = cluster_info['metadata'][:3]  # First 3 items
            for i, meta in enumerate(sample_metadata):
                logger.info(f"  Sample {i+1}: {meta.get('topic', 'N/A')} - {meta.get('summary', 'N/A')[:50]}...")
    
    # Performance benchmarking
    logger.info("\n" + "="*50)
    logger.info("PERFORMANCE BENCHMARKING")
    logger.info("="*50)
    
    if queries and len(query_embeddings) > 0:
        import time
        
        # Benchmark different search configurations
        benchmark_configs = [
            {'num_coarse_clusters': 1, 'k': 5},
            {'num_coarse_clusters': 2, 'k': 5},
            {'num_coarse_clusters': 3, 'k': 5},
            {'num_coarse_clusters': 4, 'k': 5},
        ]
        
        for config_test in benchmark_configs:
            times = []
            for query_emb in query_embeddings[:10]:  # Test on first 10 queries
                start_time = time.time()
                faiss_index.search(query_emb, **config_test)
                times.append(time.time() - start_time)
            
            avg_time = np.mean(times)
            std_time = np.std(times)
            logger.info(f"Config {config_test}: Avg time: {avg_time:.4f}s ± {std_time:.4f}s")
    
    logger.info("\nRAG model implementation with similarity-based metadata clustering completed successfully!")
    logger.info(f"Index saved to: {config.INDEX_PATH}")
    
    return {
        'faiss_index': faiss_index,
        'content_model': content_embedding_model,
        'metadata_model': metadata_embedding_model,
        'chunks': all_chunks,
        'metadata': all_metadata,
        'metrics': metrics if queries else None
    }# config.py
import os
from dataclasses import dataclass

@dataclass
class Config:
    # Model settings
    EMBEDDING_MODEL_NAME: str = "all-MiniLM-L6-v2"
    EMBEDDING_DIM: int = 384
    BATCH_SIZE: int = 64
    
    # LLM API settings
    LLM_API_URL: str = os.getenv("LLM_API_URL", "https://api.openai.com/v1/chat/completions")
    LLM_API_KEY: str = os.getenv("LLM_API_KEY", "")
    LLM_MODEL: str = "gpt-3.5-turbo"
    
    # FAISS settings
    NLIST: int = 100  # Number of clusters for metadata
    NPROBE: int = 10  # Number of clusters to search
    M: int = 8        # Number of subquantizers for PQ
    NBITS: int = 8    # Number of bits per subquantizer
    
    # Chunking settings
    MAX_CHUNK_SIZE: int = 1000
    CHUNK_OVERLAP: int = 200
    
    # File paths
    LECTURE_FOLDER: str = "lectures"
    QUERY_FILE: str = "queries.txt"
    INDEX_PATH: str = "faiss_index"

config = Config()

# requirements.txt content (add this as a comment for reference)
"""
sentence-transformers>=2.2.0
faiss-cpu>=1.7.0
numpy>=1.21.0
torch>=1.9.0
tqdm>=4.62.0
python-docx>=0.8.11
requests>=2.28.0
matplotlib>=3.5.0
scikit-learn>=1.0.0
mammoth>=1.4.0
"""

# data_processing.py
import logging
import os
import re
from typing import List, Dict, Tuple
from pathlib import Path
import mammoth
import docx
from docx import Document
from docx.table import Table

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file preserving table structure"""
        try:
            doc = Document(file_path)
            full_text = []
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = next((p for p in doc.paragraphs if p._element is element), None)
                    if para and para.text.strip():
                        full_text.append(para.text.strip())
                
                elif element.tag.endswith('tbl'):  # Table
                    table = next((t for t in doc.tables if t._element is element), None)
                    if table:
                        table_text = self.extract_table_text(table)
                        if table_text:
                            full_text.append(f"[TABLE]\n{table_text}\n[/TABLE]")
            
            return '\n\n'.join(full_text)
        
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            return ""
    
    def extract_table_text(self, table: Table) -> str:
        """Extract text from table maintaining structure"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            if any(cell for cell in row_data):  # Skip empty rows
                table_data.append(' | '.join(row_data))
        
        return '\n'.join(table_data)

def load_and_preprocess_lectures(lecture_folder: str) -> List[Dict[str, str]]:
    """Load and preprocess lecture documents from folder"""
    logger.info(f"Loading lectures from {lecture_folder}")
    
    processor = DocumentProcessor()
    documents = []
    
    lecture_path = Path(lecture_folder)
    if not lecture_path.exists():
        logger.error(f"Lecture folder {lecture_folder} does not exist")
        return []
    
    for file_path in lecture_path.rglob("*.docx"):
        try:
            logger.info(f"Processing {file_path}")
            text = processor.extract_text_from_docx(str(file_path))
            
            if text:
                documents.append({
                    'filename': file_path.name,
                    'filepath': str(file_path),
                    'content': text,
                    'file_type': 'docx'
                })
                logger.info(f"Successfully processed {file_path.name}")
            else:
                logger.warning(f"No text extracted from {file_path}")
                
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
    
    logger.info(f"Loaded {len(documents)} documents")
    return documents

def load_and_preprocess_queries(query_file: str) -> List[str]:
    """Load queries from text file"""
    logger.info(f"Loading queries from {query_file}")
    
    try:
        with open(query_file, 'r', encoding='utf-8') as f:
            queries = [line.strip() for line in f if line.strip()]
        
        logger.info(f"Loaded {len(queries)} queries")
        return queries
    
    except FileNotFoundError:
        logger.error(f"Query file {query_file} not found")
        return []
    except Exception as e:
        logger.error(f"Error loading queries: {e}")
        return []

# llm_client.py
import requests
import json
import logging
from typing import List, Dict, Optional
import time
from config import config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LLMClient:
    def __init__(self, api_url: str = None, api_key: str = None, model: str = None):
        self.api_url = api_url or config.LLM_API_URL
        self.api_key = api_key or config.LLM_API_KEY
        self.model = model or config.LLM_MODEL
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    
    def call_llm(self, prompt: str, max_retries: int = 3) -> Optional[str]:
        """Make API call to LLM with retry logic"""
        payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 1000,
            "temperature": 0.1
        }
        
        for attempt in range(max_retries):
            try:
                response = requests.post(
                    self.api_url, 
                    headers=self.headers, 
                    json=payload,
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    return result['choices'][0]['message']['content'].strip()
                else:
                    logger.error(f"API call failed with status {response.status_code}: {response.text}")
                    
            except Exception as e:
                logger.error(f"API call attempt {attempt + 1} failed: {e}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
        
        return None
    
    def intelligent_chunk(self, text: str, max_chunk_size: int = 1000) -> List[str]:
        """Use LLM to intelligently chunk text"""
        prompt = f"""
        Please divide the following text into logical chunks that preserve semantic meaning. 
        Each chunk should be approximately {max_chunk_size} characters or less.
        Focus on maintaining context and coherence within each chunk.
        Return only the chunks separated by "---CHUNK---":

        Text: {text[:4000]}  # Limit text size for API
        """
        
        response = self.call_llm(prompt)
        if response:
            chunks = [chunk.strip() for chunk in response.split("---CHUNK---") if chunk.strip()]
            return chunks if chunks else [text[:max_chunk_size]]
        else:
            # Fallback to simple chunking
            return [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
    
    def extract_metadata(self, chunk: str) -> Dict[str, str]:
        """Extract metadata from chunk using LLM"""
        prompt = f"""
        Analyze the following text chunk and extract key metadata as JSON:
        - topic: main topic/subject
        - category: document type (lecture, tutorial, reference, etc.)
        - difficulty: beginner/intermediate/advanced
        - keywords: comma-separated key terms
        - summary: brief 1-sentence summary

        Text: {chunk[:1500]}

        Return only valid JSON format.
        """
        
        response = self.call_llm(prompt)
        if response:
            try:
                # Extract JSON from response
                json_start = response.find('{')
                json_end = response.rfind('}') + 1
                if json_start != -1 and json_end != 0:
                    metadata = json.loads(response[json_start:json_end])
                    return metadata
            except json.JSONDecodeError:
                logger.warning("Failed to parse LLM metadata response as JSON")
        
        # Fallback metadata
        return {
            "topic": "general",
            "category": "document",
            "difficulty": "intermediate",
            "keywords": "",
            "summary": chunk[:100] + "..."
        }

# embeddings.py
import logging
from typing import List
from sentence_transformers import SentenceTransformer
import numpy as np
import torch
from tqdm.auto import tqdm

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def load_embedding_model(model_name: str = "all-MiniLM-L6-v2") -> SentenceTransformer:
    """Load sentence transformer model"""
    logger.info(f"Loading embedding model: {model_name}")
    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    model = SentenceTransformer(model_name, device=device)
    logger.info(f"Model loaded on device: {device}")
    return model

def compute_embeddings(text_list: List[str], model: SentenceTransformer, batch_size: int = 64) -> np.ndarray:
    """Compute embeddings for list of texts"""
    logger.info(f"Computing embeddings for {len(text_list)} texts")
    
    embeddings = []
    for i in tqdm(range(0, len(text_list), batch_size), desc="Computing embeddings"):
        batch = text_list[i:i+batch_size]
        batch_embeddings = model.encode(batch, convert_to_numpy=True, show_progress_bar=False)
        embeddings.append(batch_embeddings)
    
    embeddings = np.vstack(embeddings)
    logger.info(f"Computed embeddings shape: {embeddings.shape}")
    return embeddings

def normalize_embeddings(embeddings: np.ndarray) -> np.ndarray:
    """Normalize embeddings to unit vectors"""
    logger.info("Normalizing embeddings...")
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
    norms[norms == 0] = 1  # Avoid division by zero
    return embeddings / norms

# faiss_index.py
import logging
import numpy as np
import faiss
import pickle
from typing import List, Dict, Tuple, Optional
from collections import defaultdict
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from sklearn.cluster import KMeans
from sentence_transformers import SentenceTransformer
import threading
from config import config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class HierarchicalFAISSIndex:
    def __init__(self, embedding_dim: int = 384, metadata_embedding_model: SentenceTransformer = None):
        self.embedding_dim = embedding_dim
        self.metadata_embedding_model = metadata_embedding_model
        
        # Hierarchical structure
        self.coarse_clusters = {}           # coarse_cluster_id -> metadata
        self.coarse_cluster_embeddings = None  # embeddings of coarse cluster centroids
        self.coarse_cluster_index = None    # FAISS index for coarse clusters
        
        self.fine_cluster_indices = {}      # coarse_cluster_id -> fine FAISS indices
        self.fine_cluster_metadata = {}     # coarse_cluster_id -> list of chunk metadata
        self.fine_cluster_embeddings = {}   # coarse_cluster_id -> chunk embeddings
        
        self.global_to_local = {}           # global_id -> (coarse_id, fine_id)
        self.next_global_id = 0
        self.lock = threading.Lock()
        
    def _create_metadata_embedding(self, metadata: Dict[str, str]) -> str:
        """Create a text representation of metadata for embedding"""
        text_parts = []
        
        # Prioritize key fields
        if metadata.get('topic'):
            text_parts.append(f"Topic: {metadata['topic']}")
        if metadata.get('category'):
            text_parts.append(f"Category: {metadata['category']}")
        if metadata.get('difficulty'):
            text_parts.append(f"Difficulty: {metadata['difficulty']}")
        if metadata.get('keywords'):
            text_parts.append(f"Keywords: {metadata['keywords']}")
        if metadata.get('summary'):
            text_parts.append(f"Summary: {metadata['summary']}")
            
        return " | ".join(text_parts)
    
    def _build_coarse_clusters(self, metadata_list: List[Dict[str, str]], num_coarse_clusters: int = None):
        """Build coarse clusters based on metadata similarity"""
        logger.info("Building coarse clusters from metadata...")
        
        # Create metadata embeddings
        metadata_texts = [self._create_metadata_embedding(meta) for meta in metadata_list]
        
        if self.metadata_embedding_model:
            metadata_embeddings = self.metadata_embedding_model.encode(
                metadata_texts, convert_to_numpy=True, show_progress_bar=True
            )
        else:
            # Fallback: use simple text-based clustering
            logger.warning("No metadata embedding model provided, using fallback clustering")
            return self._fallback_clustering(metadata_list)
        
        # Normalize metadata embeddings
        metadata_embeddings = metadata_embeddings / np.linalg.norm(
            metadata_embeddings, axis=1, keepdims=True
        )
        
        # Determine optimal number of coarse clusters
        if num_coarse_clusters is None:
            num_coarse_clusters = min(max(len(metadata_list) // 20, 3), 15)
        
        # Perform K-means clustering on metadata embeddings
        logger.info(f"Performing K-means clustering with {num_coarse_clusters} clusters")
        kmeans = KMeans(n_clusters=num_coarse_clusters, random_state=42, n_init=10)
        cluster_labels = kmeans.fit_predict(metadata_embeddings)
        
        # Create coarse cluster structure
        coarse_cluster_data = defaultdict(list)
        for i, label in enumerate(cluster_labels):
            coarse_cluster_data[f"coarse_{label}"].append(i)
        
        # Store coarse cluster centroids and create FAISS index
        coarse_centroids = kmeans.cluster_centers_
        self.coarse_cluster_embeddings = coarse_centroids
        
        # Create FAISS index for coarse clusters
        self.coarse_cluster_index = faiss.IndexFlatIP(coarse_centroids.shape[1])
        self.coarse_cluster_index.add(coarse_centroids.astype(np.float32))
        
        # Store coarse cluster metadata
        for cluster_id, indices in coarse_cluster_data.items():
            cluster_metadata = [metadata_list[i] for i in indices]
            self.coarse_clusters[cluster_id] = {
                'indices': indices,
                'metadata': cluster_metadata,
                'centroid_text': self._create_cluster_summary(cluster_metadata)
            }
        
        logger.info(f"Created {len(coarse_cluster_data)} coarse clusters")
        return coarse_cluster_data
    
    def _fallback_clustering(self, metadata_list: List[Dict[str, str]]):
        """Fallback clustering when no embedding model is available"""
        clusters = defaultdict(list)
        for i, metadata in enumerate(metadata_list):
            cluster_key = f"coarse_{metadata.get('topic', 'general')}_{metadata.get('category', 'doc')}"
            clusters[cluster_key].append(i)
        return clusters
    
    def _create_cluster_summary(self, cluster_metadata: List[Dict[str, str]]) -> str:
        """Create a summary text for a cluster"""
        topics = [m.get('topic', '') for m in cluster_metadata]
        categories = [m.get('category', '') for m in cluster_metadata]
        
        # Get most common topics and categories
        from collections import Counter
        top_topics = Counter(topics).most_common(3)
        top_categories = Counter(categories).most_common(2)
        
        summary_parts = []
        if top_topics:
            summary_parts.append(f"Topics: {', '.join([t[0] for t in top_topics if t[0]])}")
        if top_categories:
            summary_parts.append(f"Categories: {', '.join([c[0] for c in top_categories if c[0]])}")
            
        return " | ".join(summary_parts)
        
    def build_hierarchical_index(self, embeddings: np.ndarray, metadata_list: List[Dict[str, str]]):
        """Build hierarchical FAISS index with metadata similarity clustering"""
        logger.info("Building hierarchical FAISS index with metadata similarity...")
        
        # Step 1: Build coarse clusters based on metadata similarity
        coarse_cluster_data = self._build_coarse_clusters(metadata_list)
        
        # Step 2: Build fine-grained FAISS indices within each coarse cluster
        logger.info("Building fine-grained indices within coarse clusters...")
        
        for coarse_cluster_id, cluster_info in self.coarse_clusters.items():
            indices = cluster_info['indices']
            
            if len(indices) < 2:  # Skip clusters with too few items
                continue
                
            cluster_embeddings = embeddings[indices]
            cluster_metadata = cluster_info['metadata']
            
            # Create appropriate FAISS index based on cluster size
            if len(indices) > 100:
                # Use IVF for larger clusters
                quantizer = faiss.IndexFlatIP(self.embedding_dim)
                nlist = min(len(indices) // 15, 30)  # Adaptive nlist
                index = faiss.IndexIVFFlat(quantizer, self.embedding_dim, nlist)
                index.train(cluster_embeddings.astype(np.float32))
            else:
                # Use flat index for smaller clusters
                index = faiss.IndexFlatIP(self.embedding_dim)
            
            index.add(cluster_embeddings.astype(np.float32))
            
            # Store fine cluster data
            self.fine_cluster_indices[coarse_cluster_id] = index
            self.fine_cluster_metadata[coarse_cluster_id] = cluster_metadata
            self.fine_cluster_embeddings[coarse_cluster_id] = cluster_embeddings
            
            # Update global to local mapping
            for local_id, global_idx in enumerate(indices):
                self.global_to_local[self.next_global_id] = (coarse_cluster_id, local_id)
                self.next_global_id += 1
        
        logger.info(f"Built hierarchical index with {len(self.fine_cluster_indices)} fine clusters")
    
    def _find_similar_coarse_clusters(self, query_embedding: np.ndarray, num_clusters: int = 3) -> List[str]:
        """Find most similar coarse clusters using metadata embedding"""
        if self.coarse_cluster_index is None or self.metadata_embedding_model is None:
            # Fallback: return all clusters
            return list(self.coarse_clusters.keys())
        
        # For content-based query, we need to map it to metadata space
        # This is a simplified approach - in practice, you might want to use query expansion
        query_text = "General query content"  # Placeholder - could be enhanced
        query_metadata_embedding = self.metadata_embedding_model.encode([query_text], convert_to_numpy=True)
        query_metadata_embedding = query_metadata_embedding / np.linalg.norm(query_metadata_embedding, axis=1, keepdims=True)
        
        # Search coarse cluster index
        scores, indices = self.coarse_cluster_index.search(
            query_metadata_embedding.astype(np.float32), 
            min(num_clusters, len(self.coarse_clusters))
        )
        
        # Return cluster IDs
        selected_clusters = []
        for idx in indices[0]:
            if idx != -1:
                cluster_id = f"coarse_{idx}"
                if cluster_id in self.coarse_clusters:
                    selected_clusters.append(cluster_id)
        
        return selected_clusters if selected_clusters else list(self.coarse_clusters.keys())[:num_clusters]
    
    def _search_cluster_parallel(self, cluster_id: str, query_embedding: np.ndarray, k: int) -> List[Tuple[float, str, int, Dict]]:
        """Search a single cluster (for parallel execution)"""
        if cluster_id not in self.fine_cluster_indices:
            return []
        
        index = self.fine_cluster_indices[cluster_id]
        metadata_list = self.fine_cluster_metadata[cluster_id]
        
        # Set nprobe for IVF indices
        if hasattr(index, 'nprobe'):
            index.nprobe = min(config.NPROBE, getattr(index, 'nlist', 10))
        
        try:
            # Search within this cluster
            scores, indices = index.search(
                query_embedding.reshape(1, -1).astype(np.float32), 
                min(k, len(metadata_list))
            )
            
            # Collect results with cluster info
            results = []
            for score, idx in zip(scores[0], indices[0]):
                if idx != -1 and idx < len(metadata_list):
                    results.append((float(score), cluster_id, idx, metadata_list[idx]))
            
            return results
        except Exception as e:
            logger.error(f"Error searching cluster {cluster_id}: {e}")
            return []
    
    def search(self, query_embedding: np.ndarray, k: int = 5, num_coarse_clusters: int = 3) -> List[Tuple[float, int, Dict]]:
        """Search across hierarchical index with parallel coarse cluster search"""
        # Step 1: Find most similar coarse clusters
        selected_clusters = self._find_similar_coarse_clusters(query_embedding, num_coarse_clusters)
        logger.debug(f"Selected coarse clusters: {selected_clusters}")
        
        # Step 2: Search selected clusters in parallel
        all_results = []
        
        with ThreadPoolExecutor(max_workers=min(len(selected_clusters), 4)) as executor:
            # Submit parallel search tasks
            future_to_cluster = {
                executor.submit(self._search_cluster_parallel, cluster_id, query_embedding, k * 2): cluster_id
                for cluster_id in selected_clusters
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_cluster):
                cluster_id = future_to_cluster[future]
                try:
                    cluster_results = future.result(timeout=10)  # 10 second timeout
                    all_results.extend(cluster_results)
                except Exception as e:
                    logger.error(f"Error in parallel search for cluster {cluster_id}: {e}")
        
        # Step 3: Merge and rank results
        # Sort by score (descending) and return top k
        all_results.sort(key=lambda x: x[0], reverse=True)
        
        # Convert to expected format (score, global_idx, metadata)
        final_results = []
        for score, cluster_id, local_idx, metadata in all_results[:k]:
            # Create a pseudo global index for compatibility
            global_idx = hash(f"{cluster_id}_{local_idx}") % 1000000
            final_results.append((score, global_idx, metadata))
        
        return final_results
    
    def save_index(self, path: str):
        """Save hierarchical index to disk"""
        os.makedirs(path, exist_ok=True)
        
        # Save coarse cluster index
        if self.coarse_cluster_index is not None:
            faiss.write_index(self.coarse_cluster_index, os.path.join(path, "coarse_index.faiss"))
        
        # Save fine cluster FAISS indices
        for cluster_id, index in self.fine_cluster_indices.items():
            safe_name = cluster_id.replace('/', '_').replace('\\', '_')
            faiss.write_index(index, os.path.join(path, f"fine_{safe_name}.index"))
        
        # Save all metadata and mappings
        metadata_path = os.path.join(path, "hierarchical_metadata.pkl")
        with open(metadata_path, 'wb') as f:
            pickle.dump({
                'coarse_clusters': self.coarse_clusters,
                'fine_cluster_metadata': self.fine_cluster_metadata,
                'coarse_cluster_embeddings': self.coarse_cluster_embeddings,
                'global_to_local': self.global_to_local,
                'embedding_dim': self.embedding_dim,
                'next_global_id': self.next_global_id
            }, f)
        
        logger.info(f"Saved hierarchical index to {path}")
    
    def load_index(self, path: str):
        """Load hierarchical index from disk"""
        # Load metadata and mappings
        metadata_path = os.path.join(path, "hierarchical_metadata.pkl")
        with open(metadata_path, 'rb') as f:
            data = pickle.load(f)
            self.coarse_clusters = data.get('coarse_clusters', {})
            self.fine_cluster_metadata = data.get('fine_cluster_metadata', {})
            self.coarse_cluster_embeddings = data.get('coarse_cluster_embeddings', None)
            self.global_to_local = data.get('global_to_local', {})
            self.embedding_dim = data.get('embedding_dim', 384)
            self.next_global_id = data.get('next_global_id', 0)
        
        # Load coarse cluster index
        coarse_index_path = os.path.join(path, "coarse_index.faiss")
        if os.path.exists(coarse_index_path):
            self.coarse_cluster_index = faiss.read_index(coarse_index_path)
        
        # Load fine cluster FAISS indices
        self.fine_cluster_indices = {}
        for cluster_id in self.fine_cluster_metadata.keys():
            safe_name = cluster_id.replace('/', '_').replace('\\', '_')
            index_path = os.path.join(path, f"fine_{safe_name}.index")
            if os.path.exists(index_path):
                self.fine_cluster_indices[cluster_id] = faiss.read_index(index_path)
        
        logger.info(f"Loaded hierarchical index from {path}")

def evaluate_faiss_index(index: HierarchicalFAISSIndex, query_embeddings: np.ndarray, 
                        ground_truth: List[List[int]] = None) -> Dict[str, float]:
    """Evaluate FAISS index performance"""
    logger.info("Evaluating hierarchical FAISS index...")
    
    total_queries = len(query_embeddings)
    avg_recall = 0.0
    search_times = []
    cluster_distribution = defaultdict(int)
    
    import time
    for i, query_emb in enumerate(query_embeddings):
        start_time = time.time()
        results = index.search(query_emb, k=10, num_coarse_clusters=3)
        search_time = time.time() - start_time
        search_times.append(search_time)
        
        # Track cluster distribution
        for score, idx, metadata in results:
            cluster_key = f"{metadata.get('topic', 'unknown')}_{metadata.get('category', 'unknown')}"
            cluster_distribution[cluster_key] += 1
        
        if ground_truth and i < len(ground_truth):
            retrieved_ids = [r[1] for r in results]
            relevant_ids = ground_truth[i]
            if relevant_ids:
                recall = len(set(retrieved_ids) & set(relevant_ids)) / len(relevant_ids)
                avg_recall += recall
    
    metrics = {
        'avg_search_time': np.mean(search_times),
        'std_search_time': np.std(search_times),
        'avg_recall': avg_recall / total_queries if ground_truth else 0.0,
        'total_coarse_clusters': len(index.coarse_clusters),
        'total_fine_clusters': len(index.fine_cluster_indices),
        'cluster_distribution': dict(cluster_distribution)
    }
    
    logger.info(f"Evaluation metrics: {metrics}")
    return metrics

# main.py
import logging
import numpy as np
import matplotlib.pyplot as plt
from data_processing import load_and_preprocess_lectures, load_and_preprocess_queries
from embeddings import load_embedding_model, compute_embeddings, normalize_embeddings
from faiss_index import HierarchicalFAISSIndex, evaluate_faiss_index
from llm_client import LLMClient
from config import config
import os

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def main():
    logger.info("Starting RAG model implementation...")
    
    # Initialize components
    llm_client = LLMClient()
    embedding_model = load_embedding_model(config.EMBEDDING_MODEL_NAME)
    faiss_index = HierarchicalFAISSIndex(config.EMBEDDING_DIM)
    
    # Load and preprocess documents
    logger.info("Loading and preprocessing documents...")
    documents = load_and_preprocess_lectures(config.LECTURE_FOLDER)
    
    if not documents:
        logger.error("No documents loaded. Please check the lecture folder.")
        return
    
    # Process each document with LLM chunking and metadata extraction
    all_chunks = []
    all_metadata = []
    
    for doc in documents:
        logger.info(f"Processing document: {doc['filename']}")
        
        # Intelligent chunking using LLM
        chunks = llm_client.intelligent_chunk(doc['content'], config.MAX_CHUNK_SIZE)
        
        for chunk in chunks:
            # Extract metadata using LLM
            metadata = llm_client.extract_metadata(chunk)
            metadata['source_file'] = doc['filename']
            metadata['chunk_text'] = chunk
            
            all_chunks.append(chunk)
            all_metadata.append(metadata)
        
        logger.info(f"Created {len(chunks)} chunks from {doc['filename']}")
    
    logger.info(f"Total chunks created: {len(all_chunks)}")
    
    # Compute embeddings for chunks
    logger.info("Computing embeddings for chunks...")
    chunk_embeddings = compute_embeddings(all_chunks, embedding_model, config.BATCH_SIZE)
    chunk_embeddings = normalize_embeddings(chunk_embeddings)
    
    # Build hierarchical FAISS index
    logger.info("Building hierarchical FAISS index...")
    faiss_index.build_hierarchical_index(chunk_embeddings, all_metadata)
    
    # Save the index
    if not os.path.exists(config.INDEX_PATH):
        os.makedirs(config.INDEX_PATH)
    faiss_index.save_index(config.INDEX_PATH)
    
    # Load and process queries
    queries = load_and_preprocess_queries(config.QUERY_FILE)
    
    if queries:
        logger.info("Processing queries...")
        query_embeddings = compute_embeddings(queries, embedding_model, config.BATCH_SIZE)
        query_embeddings = normalize_embeddings(query_embeddings)
        
        # Evaluate the index
        metrics = evaluate_faiss_index(faiss_index, query_embeddings)
        
        # Test search functionality
        logger.info("Testing search functionality...")
        for i, (query, query_emb) in enumerate(zip(queries[:3], query_embeddings[:3])):
            logger.info(f"\nQuery {i+1}: {query}")
            results = faiss_index.search(query_emb, k=5)
            
            logger.info("Top 5 results:")
            for j, (score, idx, metadata) in enumerate(results):
                logger.info(f"  {j+1}. Score: {score:.4f}")
                logger.info(f"     Topic: {metadata['topic']}")
                logger.info(f"     Summary: {metadata['summary']}")
                logger.info(f"     Source: {metadata['source_file']}")
    
    logger.info("RAG model implementation completed successfully!")

if __name__ == '__main__':
    main()